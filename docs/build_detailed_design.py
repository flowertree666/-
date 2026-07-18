from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from datetime import date
from pathlib import Path

OUT = Path(r"D:\18\详细设计文档.docx")

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "18342D"
MUTED = "63736D"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GREEN = "2F6F59"
GOLD = "A56A22"
RED = "9B1C1C"
WHITE = "FFFFFF"
GRID = "C9D3DC"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_run_font(run, name="Microsoft YaHei", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def set_para_format(p, before=0, after=6, line=1.25, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
]:
    s = styles[style_name]
    s.font.name = "Microsoft YaHei"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

if "Code Block" not in styles:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code Block"]
code_style.font.name = "Consolas"
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor.from_string("253238")
code_style.paragraph_format.left_indent = Inches(0.25)
code_style.paragraph_format.right_indent = Inches(0.15)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(8)
code_style.paragraph_format.line_spacing = 1.05

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_format(hp, after=0, line=1.0)
r = hp.add_run("智能旅游助手  |  详细设计文档")
set_run_font(r, size=9, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_format(fp, after=0, line=1.0)
r = fp.add_run("内部技术文档  |  第 ")
set_run_font(r, size=9, color=MUTED)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)
r = fp.add_run(" 页")
set_run_font(r, size=9, color=MUTED)

def add_title(text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, size=28, bold=True, color=INK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p2, after=30, line=1.1)
        r2 = p2.add_run(subtitle)
        set_run_font(r2, size=14, color=MUTED)

def add_h1(text, page_break=False):
    p = doc.add_paragraph(text, style="Heading 1")
    return p

def add_h2(text):
    return doc.add_paragraph(text, style="Heading 2")

def add_h3(text):
    return doc.add_paragraph(text, style="Heading 3")

def add_p(text, bold_prefix=None, color=None):
    p = doc.add_paragraph()
    set_para_format(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color or INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=color or INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color or INK)
    return p

def add_bullets(items, level=0):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "•")
    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(540 + 360 * level))
    ind.set(qn("w:hanging"), "270")
    pPr.append(ind)
    lvl.extend([start, numFmt, lvlText, pPr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(num_id))
        numPr.extend([ilvl, numId])
        p._p.get_or_add_pPr().append(numPr)
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)

def add_numbers(items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "decimal")
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    pPr.append(tabs)
    pPr.append(ind)
    lvl.extend([start, numFmt, lvlText, suff, pPr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(num_id))
        numPr.extend([ilvl, numId])
        p._p.get_or_add_pPr().append(numPr)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(widths[idx]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, after=0, line=1.1)
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(widths) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_para_format(p, after=0, line=1.15)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table

def add_callout(label, text, fill=PALE, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_format(p, after=0, line=1.2)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_code(text):
    p = doc.add_paragraph(style="Code Block")
    set_cell = OxmlElement("w:shd")
    set_cell.set(qn("w:fill"), "F3F5F7")
    p._p.get_or_add_pPr().append(set_cell)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=9, color="253238")

# Cover
add_title("智能旅游助手详细设计文档", "基于 Vue 3、Spring Boot、MySQL 与通用智能体能力")
add_callout("文档定位", "面向课程设计、作品答辩、开发维护与部署验收，描述当前项目的实际实现、关键算法、接口契约、数据模型和运行边界。", fill="EAF4EF", accent=GREEN)
cover_meta = [
    ("项目名称", "基于通用智能体的智能旅游助手"),
    ("技术栈", "Vue 3 + Vite / Java 21 + Spring Boot / MySQL"),
    ("外部能力", "高德地图、博查搜索、DeepSeek、百度实时路况、OpenWeather/WeatherAPI"),
    ("文档版本", "V1.0"),
    ("编制日期", "2026-07-16"),
    ("文档状态", "详细设计基线"),
]
add_table(["项目", "内容"], cover_meta, [2700, 6660], 10)
doc.add_page_break()

add_h1("修订记录")
add_table(["版本", "日期", "修订内容", "状态"], [
    ("V1.0", "2026-07-16", "依据当前 18-4-test 后续代码形成详细设计基线", "已完成"),
], [1200, 1700, 5160, 1300])

add_h1("目录")
toc_items = [
    "1. 文档概述", "2. 项目背景与目标", "3. 需求分析", "4. 总体架构设计", "5. 前端详细设计",
    "6. 后端详细设计", "7. 联网研究与证据链设计", "8. 路线规划引擎设计", "9. 天气与客流预测设计",
    "10. 实时交通压力采样设计", "11. 预算与住宿餐饮设计", "12. 数据库设计", "13. 接口设计",
    "14. 异常处理与降级策略", "15. 安全与合规设计", "16. 部署与运行", "17. 测试设计", "18. 性能与扩展设计",
]
add_bullets(toc_items)
add_callout("阅读建议", "第 4 章说明系统全貌；第 8、10、11 章是核心业务设计；第 12、13 章可直接用于数据库和接口答辩；第 16、17 章用于部署与验收。")

add_h1("1. 文档概述", page_break=True)
add_h2("1.1 编写目的")
add_p("本文档定义智能旅游助手的详细技术设计，作为编码实现、接口联调、数据库建设、测试验收、作品演示和后续扩展的共同依据。文档描述的是当前工作区代码已经形成的能力，并对生产环境配置给出明确边界。")
add_h2("1.2 适用范围")
add_bullets([
    "用户输入中国境内城市、区县、旅行天数、日期、偏好、预算和自然语言特殊要求。",
    "系统联网获取近三个月旅游信息，生成带来源网址的研究报告。",
    "系统输出经济、标准、舒适三套差异化路线，并生成每日时间线、地图路线、餐饮、住宿、天气、预算和交通压力信息。",
    "系统保存研究快照、行程历史和交通采样数据；默认演示环境为 H2，生产环境使用 MySQL。",
])
add_h2("1.3 术语")
add_table(["术语", "定义"], [
    ("POI", "Point of Interest，景点、餐厅、酒店等地理兴趣点。"),
    ("GCJ-02", "高德地图采用的中国加密坐标系。"),
    ("客流预测", "由景点热度、空间引力、时段、天气等特征计算的参考指标，不是园内实时人数。"),
    ("交通压力", "百度实时路况反映的景区周边机动车道路状态，不等同于园内游客密度。"),
    ("研究快照", "某次联网搜索和 AI 分析的完整 JSON 结果，包含来源、日期和热度排行。"),
], [1800, 7560])

add_h1("2. 项目背景与目标", page_break=True)
add_h2("2.1 背景")
add_p("传统旅游攻略往往存在信息分散、更新滞后、路线模板化和人工筛选成本高的问题。本项目以通用智能体为调度核心，将联网研究、地图 POI、路线规划、天气分析、预算估算和历史记录整合为一条可解释的规划链路。")
add_h2("2.2 建设目标")
add_bullets([
    "全国化：输入真实城市或区县即可动态同步景点，而非依赖单一城市静态数据。",
    "个性化：特殊期待中的指定地点、指定天数、登山、日出和起床时间形成硬约束。",
    "可落地：时间线包含起床、三餐、景点、交通、入住、退房和行李处理。",
    "可解释：热度结论必须标注具体来源 URL；预测与实时数据必须区分。",
    "多方案：围绕预算生成经济、标准、舒适三套内容差异明显的组合。",
    "可追溯：保存行程、研究报告和五分钟交通采样快照。",
])
add_h2("2.3 非目标与边界")
add_callout("数据边界", "系统不绕过登录、验证码、签名或反爬机制采集抖音、小红书等受限数据；不伪造评论正文、票价、园内实时人数或平台商户 ID。", fill="FFF6E5", accent=GOLD)

add_h1("3. 需求分析", page_break=True)
add_h2("3.1 功能需求")
add_table(["编号", "功能", "详细说明"], [
    ("FR-01", "目的地校验", "通过高德地理编码校验城市或区县，不存在或过短的输入直接报错。"),
    ("FR-02", "用户画像", "解析天数、日期、偏好、交通方式、预算和自由文本。"),
    ("FR-03", "联网研究", "搜索目标日期前三个月内的具体页面，保留发布日、正文摘要和 URL。"),
    ("FR-04", "多路线规划", "生成经济、标准、舒适三套方案并按天拆分。"),
    ("FR-05", "时间线", "动态计算起床、早餐、交通、景点、午餐、晚餐和住宿节点。"),
    ("FR-06", "地图路线", "按天独立绘制步行或驾车真实道路，支持缩放和日切换。"),
    ("FR-07", "住宿行李", "按空间跨度决定是否换酒店，遵守 13:00 入住/退房约束并规划寄存。"),
    ("FR-08", "预算", "以用户预算为基准计算 70%、100%、130% 三档目标及每日费用。"),
    ("FR-09", "天气", "查询目标日期天气并给出降雨、高温和风力建议。"),
    ("FR-10", "交通压力", "景点周边路况立即采样、五分钟更新、噪声标记并进行历史预测。"),
    ("FR-11", "历史记录", "保存并查看行程及当次研究快照，支持删除。"),
], [900, 1800, 6660], 8.8)
add_h2("3.2 非功能需求")
add_table(["类别", "要求"], [
    ("可用性", "外部 API 失败时返回明确降级信息，不让整条行程生成失败。"),
    ("性能", "相同路线、餐饮、酒店和交通采样结果应缓存或复用；外部请求必须限速。"),
    ("安全", "密钥仅通过环境变量注入；后端密钥不得返回前端或写入仓库。"),
    ("数据质量", "动态数据携带来源、采集时间和置信度；预测不得包装成官方实时数据。"),
    ("可扩展性", "通过数据提供器和服务边界扩展新平台、预测模型或地图供应商。"),
], [1800, 7560])

add_h1("4. 总体架构设计", page_break=True)
add_h2("4.1 架构分层")
add_table(["层级", "主要组件", "职责"], [
    ("表现层", "Vue 3、App.vue、RouteMap.vue", "采集需求，展示研究报告、方案、时间线、地图、预算、天气和交通压力。"),
    ("接口层", "Spring MVC Controllers", "参数校验、REST 路由、异常映射和跨域访问。"),
    ("业务层", "RoutePlannerService、WebResearchService、TrafficMonitoringService", "编排搜索、意图解析、路线规划、采样和持久化。"),
    ("能力适配层", "AmapService、DeepSeekService、WeatherForecastService", "封装地图、模型、天气和外部服务差异。"),
    ("数据层", "Spring Data JPA、MySQL/H2", "保存景点、研究、行程、监测目标和路况快照。"),
    ("外部服务", "高德、百度、博查、DeepSeek、OpenWeather/WeatherAPI", "提供 POI、路线、路况、搜索、AI 和天气数据。"),
], [1300, 2700, 5360], 8.8)
add_h2("4.2 核心调用链")
add_code("用户提交 -> 目的地校验 -> 近三个月联网研究 -> AI 解析硬约束\n"
         "       -> 同步/加载景点 -> 三档路线规划 -> 餐饮住宿与交通补全\n"
         "       -> 天气与客流预测 -> 实时交通采样 -> AI 文案个性化\n"
         "       -> 保存历史快照 -> 返回前端展示")
add_h2("4.3 技术选型")
add_table(["技术", "版本/类型", "选型理由"], [
    ("Vue", "Vue 3 + Composition API", "组件状态清晰，适合单页规划交互。"),
    ("Vite", "开发与构建工具", "启动快，环境变量和 HMR 使用简单。"),
    ("Java", "Java 21", "Record、现代语法和长期维护能力。"),
    ("Spring Boot", "3.4.x", "REST、校验、JPA、调度和配置体系完整。"),
    ("MySQL", "8.x 推荐", "关系数据和 JSON 快照并存，便于查询和审计。"),
    ("H2", "内存演示库", "零配置演示；重启清空，不用于生产持久化。"),
], [1700, 2200, 5460])

add_h1("5. 前端详细设计", page_break=True)
add_h2("5.1 页面状态模型")
add_table(["状态", "类型", "用途"], [
    ("form", "Ref<Object>", "城市、天数、日期、偏好、预算和特殊期待。"),
    ("research", "Ref<ResearchReport>", "联网研究、热度排行和证据来源。"),
    ("result", "Ref<PlanResponse>", "三套规划结果。"),
    ("active / activeDay", "Ref<number>", "当前预算方案和当前天。"),
    ("histories", "Ref<Array>", "历史记录摘要。"),
    ("loading / error", "Ref", "提交进度和错误提示。"),
], [2100, 2100, 5160])
add_h2("5.2 主要交互")
add_numbers([
    "输入目的地后失焦调用 /api/amap/validate；无效目的地在表单内直接提示。",
    "提交时先调用 /api/research/analyze，再调用 /api/plans。",
    "方案切换重置 activeDay；每日地图和详细时间线只展示同一天内容。",
    "时间线将景点、餐饮、交通、住宿、休息节点合并后按 time 排序。",
    "前端每五分钟刷新交通信息；请求按景点、日期和时刻去重。",
    "餐厅链接使用城市、店名、地址构造美团搜索；住宿链接构造携程酒店搜索。",
])
add_h2("5.3 地图组件")
add_bullets([
    "RouteMap.vue 使用高德 JS API，按 activeDay 清空并重绘标记与折线。",
    "步行距离较短时调用 AMap.Walking；其他情况调用 AMap.Driving。",
    "住宿点加入每日路线，换酒店时额外绘制酒店间转场。",
    "真实路线失败时使用虚线降级，并在地图状态栏说明。",
])
add_h2("5.4 前端异常状态")
add_table(["场景", "界面行为"], [
    ("目的地无效", "显示字段错误，停止研究和规划请求。"),
    ("研究数据不足", "显示热度 0、数据不足说明和已有来源，不伪造排行。"),
    ("地图 Key 缺失", "显示地图配置提示，时间线仍可使用。"),
    ("交通无覆盖", "显示暂无道路路况，而不是无限显示正在采样。"),
    ("价格缺失", "以加粗的价格待确认标记展示。"),
], [2400, 6960])

add_h1("6. 后端详细设计", page_break=True)
add_h2("6.1 包结构")
add_table(["包", "核心类", "职责"], [
    ("planner", "RoutePlannerService", "主业务编排、路线构造、时间线和预算。"),
    ("ai", "DeepSeekService", "特殊期待解析和文案个性化。"),
    ("amap", "AmapService", "目的地、POI、路线、餐饮、住宿和基础天气。"),
    ("research", "WebResearchService", "搜索、网页校验、正文提取、证据分析。"),
    ("crowd", "CrowdPredictionService", "游览拥挤度预测。"),
    ("traffic", "TrafficMonitoringService", "实时道路采样、存储和历史预测。"),
    ("weather", "WeatherForecastService", "天气供应商主备切换。"),
    ("history", "HistoryService", "行程及研究快照保存和恢复。"),
], [1700, 3000, 4660], 8.8)
add_h2("6.2 主业务事务边界")
add_p("规划流程以一次 POST /api/plans 为业务单元。外部 API 调用不纳入数据库事务，避免长事务占用连接；研究报告和行程历史分别保存完整 JSON 快照。外部能力失败时优先使用已加载景点、规则引擎和明确的降级结果。")
add_h2("6.3 输入验证")
add_table(["字段", "约束", "处理"], [
    ("city", "非空且至少 2 个字符", "通过高德行政区地理编码进一步校验。"),
    ("days", "1-10", "超出范围返回 400。"),
    ("startDate", "非空", "作为研究窗口、天气和每日日期基准。"),
    ("budget", "500-200000", "为空时使用 3000 元默认值。"),
    ("freeText", "自由文本", "DeepSeek 解析后形成地点、日期、强度和起床时间约束。"),
], [1800, 2500, 5060])

add_h1("7. 联网研究与证据链设计", page_break=True)
add_h2("7.1 数据流程")
add_numbers([
    "按城市、目标年月、官方公告、时令活动、游客反馈生成查询词。",
    "通过博查 Web Search API 获取具体结果 URL、摘要和发布日期。",
    "过滤首页、搜索页、不安全 URL、重复 URL 和单域过量结果。",
    "遵守 robots.txt，使用 Jsoup 访问正文并剔除导航、脚本和表单。",
    "只保留目标日期前三个月内且发布日期可解析的页面。",
    "DeepSeek 仅基于成功提取的网页证据生成摘要、排行和提示。",
    "每条事实使用 [S编号] 引用，并将具体 URL 保存到 evidence。",
])
add_h2("7.2 热度评分语义")
add_callout("重要说明", "heatScore 是本次证据样本内的相对热度，不是平台官方指数，也不能等同于真实游客人数。证据不足时必须返回数据不足。", fill="FFF6E5", accent=GOLD)
add_h2("7.3 时间窗口")
add_code("windowStart = targetDate.minusMonths(3)\nwindowEnd   = min(targetDate, today)\naccept(page) = publishedDate in [windowStart, windowEnd]")
add_h2("7.4 来源质量控制")
add_table(["规则", "设计"], [
    ("具体页面", "拒绝站点首页、搜索结果页和频道列表页。"),
    ("日期可验证", "优先页面 meta/time，次选搜索 API datePublished。"),
    ("域名多样性", "同一域名默认最多保留三条，降低单源偏差。"),
    ("官方优先", "政府、文旅局和景区官网用于开放、预约和限流事实。"),
    ("失败透明", "抓取失败不补写事实，报告显示数据不足。"),
], [2100, 7260])

add_h1("8. 路线规划引擎设计", page_break=True)
add_h2("8.1 约束分类")
add_table(["约束类型", "示例", "处理方式"], [
    ("硬约束", "第二天华山全天、凌晨看日出", "必须进入指定天；不可行时返回原因。"),
    ("安全约束", "日出却要求下午出发", "拒绝生成并提示修改为清晨或日落。"),
    ("体力约束", "全天登山后的次日", "降低时长并排除登山和长距离徒步。"),
    ("空间约束", "连续三天不得局限同一区域", "选择跨区域日锚点并限制每日半径。"),
    ("时间约束", "开放时间、起床时间、交通和用餐", "游览开始前动态推进时间游标。"),
    ("软约束", "摄影、美食、自然风光", "进入景点排序权重。"),
], [1800, 2900, 4660])
add_h2("8.2 多方案差异化")
add_table(["维度", "经济", "标准", "舒适"], [
    ("目标预算", "用户预算 × 70%", "用户预算 × 100%", "用户预算 × 130%"),
    ("每日游览上限", "约 360 分钟", "约 480 分钟", "约 540 分钟"),
    ("公共交通日半径", "约 12 km", "约 20 km", "约 30 km"),
    ("排序偏好", "低成本、近距离、免费景点", "热度、距离和成本平衡", "高评分、热门、覆盖更广"),
    ("组合去重", "优先选取第一组", "避开经济方案非必选景点", "避开前两组非必选景点"),
], [1800, 2500, 2500, 2560], 8.5)
add_h2("8.3 景点评分模型")
add_code("score = preferenceMatch + seasonalFit + recentHeat + baseHeat\n"
         "        - distancePenalty - duplicatePenalty - fatiguePenalty\n"
         "requiredPlace -> highest priority; excludedPlace -> removed")
add_h2("8.4 每日时间线算法")
add_numbers([
    "从用户起床时间开始，默认 07:00；日出约束可调整至 04:00。",
    "加入 30 分钟洗漱、住宿寄存或退房、早餐和前往首个景点的真实交通。",
    "若早于开放时间，插入等待节点；游览时长来自 POI 类型或全天约束。",
    "跨越午餐窗口时查找景点附近独立餐厅，加入往返交通和 60 分钟用餐。",
    "景点间使用高德步行、公交/地铁或驾车路线推进时间游标。",
    "行程结束后办理入住、安排晚餐并返回住宿。",
])
add_h2("8.5 地铁线路提示")
add_p("公交路径解析高德 transits.segments.bus.buslines，提取线路名、上车站和下车站。若包含地铁或“X号线”，时间线显示“地铁 1号线：A站→B站；2号线：B站→C站”，最多展示三段换乘。")

add_h1("9. 天气与客流预测设计", page_break=True)
add_h2("9.1 天气供应链")
add_p("WeatherForecastService 根据景点平均坐标和目标日期查询天气。OpenWeather 为主数据源，WeatherAPI 为备用数据源；均不可用时返回明确的暂不可用信息，不阻塞路线生成。")
add_h2("9.2 客流预测模型")
add_code("gravity = Σ(otherHeat / (1 + distance)^1.35)\n"
         "spatial = 100 × (1 - exp(-gravity / 170))\n"
         "raw = (baseCrowd×0.35 + heat×0.30 + spatial×0.35)\n"
         "      × timeFactor × weekendFactor × weatherFactor")
add_table(["特征", "作用"], [
    ("景点 crowdIndex", "景点基础拥挤水平。"),
    ("heatScore", "景点基础热度和近期证据热度。"),
    ("空间引力", "周边高热度 POI 越集中，潜在客流越高。"),
    ("时段", "10-11 点、14-16 点提高系数，早 9 点前降低。"),
    ("周末", "周末提高预测值。"),
    ("天气", "降雨、高温和恶劣天气降低户外客流。"),
], [2200, 7160])
add_h2("9.3 输出")
add_p("模型输出 5-98 的指数、空闲/舒适/较拥挤/拥挤等级、推荐到达时段、建议和预测依据。前端必须同时显示模型依据与官方公告复核提示。")

add_h1("10. 实时交通压力采样设计", page_break=True)
add_h2("10.1 数据定位")
add_callout("语义隔离", "百度路况用于表示景区周边机动车交通压力，可辅助判断到达难度和潜在聚集，但不能直接证明景区园内人数。", fill="EAF4EF", accent=GREEN)
add_h2("10.2 全国通用查询链")
add_code("景点 GCJ-02 坐标\n"
         "  -> 周边 800m 路况\n"
         "  -> 有效性校验（状态/描述/道路至少一项有效）\n"
         "  -> 空结果则扩大至 1000m\n"
         "  -> 仍为空则反向地理编码最近主干道\n"
         "  -> 道路名称路况查询\n"
         "  -> 成功快照入库；临时失败保留最近成功快照")
add_h2("10.3 调度与并发")
add_bullets([
    "行程生成时注册 TrafficWatch 并立即采样。",
    "@Scheduled 每五分钟刷新未过期的监测目标。",
    "同一 attractionId 使用锁防止并发重复采样。",
    "百度请求之间至少间隔 250 ms，降低 status=302 限流风险。",
    "前端五分钟刷新，并按景点、日期和时刻去重请求。",
])
add_h2("10.4 有效性与降级")
add_table(["情况", "处理"], [
    ("HTTP/API 成功但内容为空", "判定为无效，不写入“未知成功快照”，继续下一级查询。"),
    ("步行街 road_name 错误", "自动改用坐标周边区域路况。"),
    ("大景区 800m 无道路", "扩大至 1000m；仍无数据则反查主干道。"),
    ("临时限流", "保留最近成功快照，稍后自动重试。"),
    ("百度无道路覆盖", "明确显示暂无道路路况，不伪造数据。"),
], [2700, 6660])
add_h2("10.5 噪声识别与预测")
add_bullets([
    "07:00-10:00、17:00-20:00 标记为通勤高峰噪声。",
    "周末/节假日作为独立特征保留，不与工作日直接混算。",
    "严重拥堵且速度低于 12 km/h 标记为疑似异常交通，不能直接断言发生事故。",
    "预测使用近 60 天同星期、相邻小时且排除高峰/节假日/异常样本的平均状态。",
])

add_h1("11. 预算、餐饮、住宿与行李设计", page_break=True)
add_h2("11.1 预算档位")
add_p("用户预算 B 对应经济 0.7B、标准 B、舒适 1.3B；每档允许区间为目标值 ±10%。费用按住宿、早餐、午餐、晚餐、门票、交通和弹性预留拆分。")
add_h2("11.2 POI 差异化选择")
add_bullets([
    "经济档优先较低人均、距离近、评分可接受的餐厅和酒店。",
    "标准档在评分、价格和距离之间平衡。",
    "舒适档提高评分和品质权重，并排除前两档已选门店。",
    "餐饮过滤酒店、宾馆、民宿内部餐厅，优先独立经营门店。",
])
add_h2("11.3 住宿切换")
add_p("系统计算每日景点平均中心与当前酒店距离。公共交通下超过约 8 km、驾车下超过约 18 km 时考虑换住宿；候选酒店之间距离超过约 3 km 才确认切换，避免频繁搬运行李。")
add_h2("11.4 入住与行李规则")
add_table(["场景", "规则"], [
    ("首日未到入住时间", "先到酒店前台寄存，13:00 后办理入住。"),
    ("连续入住", "行李留在客房，无需随身游览。"),
    ("更换酒店", "13:00 前退房，携带行李前往新酒店寄存，13:00 后入住。"),
    ("酒店间转场", "高德规划步行、公共交通或打车路线并计入时间线。"),
], [2500, 6860])

add_h1("12. 数据库设计", page_break=True)
add_h2("12.1 数据库策略")
add_p("生产环境使用 MySQL 8.x，字符集 utf8mb4；开发默认使用 H2 内存库并启用 MySQL 兼容模式。H2 的 ddl-auto=create-drop 会在后端重启后清空数据，正式演示历史持久化时应启用 mysql profile。")
add_h2("12.2 attraction")
add_table(["字段", "类型", "约束/说明"], [
    ("id", "BIGINT", "主键，自增。"), ("name", "VARCHAR(100)", "景点名称，非空。"),
    ("city/district", "VARCHAR", "行政区定位。"), ("category/tags", "VARCHAR", "分类与偏好标签。"),
    ("summary", "VARCHAR(1000)", "景点摘要。"), ("longitude/latitude", "DOUBLE", "GCJ-02 坐标。"),
    ("duration_minutes", "INT", "建议游览分钟。"), ("heat_score/crowd_index", "INT", "基础热度与拥挤基线。"),
    ("opening_hours", "VARCHAR", "开放时间。"), ("source_url", "VARCHAR(500)", "具体数据来源。"),
    ("collected_at/confidence", "DATETIME/DOUBLE", "采集时间与置信度。"),
], [2200, 2200, 4960], 8.7)
add_h2("12.3 travel_history")
add_table(["字段", "类型", "说明"], [
    ("id", "BIGINT", "主键。"), ("title/city/days/start_date", "基础字段", "历史列表摘要。"),
    ("preferences/pace/transport/free_text", "VARCHAR", "原始用户需求。"),
    ("plan_json", "LONGTEXT", "完整 PlanResponse 快照。"),
    ("research_json", "LONGTEXT", "当次研究报告快照。"),
    ("created_at", "DATETIME", "生成时间。"),
], [2200, 2300, 4860])
add_h2("12.4 research_insight")
add_table(["字段", "类型", "说明"], [
    ("id", "BIGINT", "主键。"), ("city", "VARCHAR(50)", "目标城市。"),
    ("report_json", "LONGTEXT", "热度、证据、排行、风险和时间戳。"),
    ("created_at", "DATETIME", "研究时间。"),
], [2200, 2300, 4860])
add_h2("12.5 traffic_watch")
add_table(["字段", "类型", "说明"], [
    ("attraction_id", "BIGINT UNIQUE", "每个景点一个活动监测目标。"),
    ("attraction_name/city", "VARCHAR", "展示和路况查询上下文。"),
    ("longitude/latitude", "DOUBLE", "采样中心。"),
    ("expires_on", "DATE", "超过行程日期后停止调度。"),
    ("updated_at", "DATETIME", "监测目标更新时间。"),
], [2200, 2300, 4860])
add_h2("12.6 traffic_snapshot")
add_table(["字段", "类型", "说明"], [
    ("attraction_id", "BIGINT", "景点关联；与 collected_at 建联合索引。"),
    ("road_name", "VARCHAR", "参考道路或周边区域。"),
    ("traffic_status", "INT", "0未知、1畅通、2缓行、3拥堵、4严重拥堵。"),
    ("description", "VARCHAR(1500)", "百度语义化路况描述。"),
    ("average_speed", "DOUBLE", "拥堵路段平均速度 km/h。"),
    ("congestion_distance/trend", "INT/VARCHAR", "拥堵距离和十分钟前趋势。"),
    ("peak_hour_noise", "BOOLEAN", "早晚高峰噪声标记。"),
    ("holiday", "BOOLEAN", "周末/节假日标记。"),
    ("possible_incident", "BOOLEAN", "疑似异常交通标记。"),
    ("collected_at", "DATETIME", "五分钟采样时间。"),
], [2400, 2300, 4660], 8.7)
add_h2("12.7 建议索引")
add_bullets([
    "attraction(city, collected_at)：按城市查询与新鲜度过滤。",
    "research_insight(city, created_at)：获取城市最新研究。",
    "travel_history(created_at)：历史列表倒序。",
    "traffic_snapshot(attraction_id, collected_at)：历史路况预测主索引。",
])

add_h1("13. 接口设计", page_break=True)
add_h2("13.1 接口总览")
add_table(["方法", "路径", "用途"], [
    ("POST", "/api/plans", "生成完整三档行程。"),
    ("GET", "/api/health", "后端健康检查。"),
    ("GET", "/api/amap/validate", "目的地合法性校验。"),
    ("POST", "/api/amap/sync", "手动同步城市景点。"),
    ("GET", "/api/amap/weather", "城市实时天气。"),
    ("GET", "/api/attractions", "按城市查询景点。"),
    ("POST", "/api/research/analyze", "生成并保存研究报告。"),
    ("GET", "/api/research/latest", "获取城市最新研究。"),
    ("GET", "/api/traffic/{id}", "景点实时交通与预测。"),
    ("GET", "/api/history", "历史摘要列表。"),
    ("GET", "/api/history/{id}", "历史详情。"),
    ("DELETE", "/api/history/{id}", "删除历史记录。"),
], [1100, 3100, 5160], 8.8)
add_h2("13.2 POST /api/plans")
add_code('{\n  "city": "杭州",\n  "days": 3,\n  "startDate": "2026-07-20",\n  "preferences": ["自然风光", "人文古迹"],\n  "pace": "适中",\n  "transport": "公共交通",\n  "budget": 3000,\n  "freeText": "第二天去华山全天徒步"\n}')
add_p("响应包含 profile、plans[3]、generatedAt、dataNotice 和 ai。每个 plan 包含 days、alternatives 和 budget；每个 day 包含 weather、accommodation、scheduleNotes、stops 和 costs。")
add_h2("13.3 GET /api/traffic/{attractionId}")
add_code("GET /api/traffic/9?date=2026-07-20&time=13:51")
add_table(["响应字段", "说明"], [
    ("liveLevel", "当前周边交通压力。"), ("roadName", "参考道路或区域。"),
    ("description", "实时语义描述。"), ("sampledAt", "精确到分钟的采样时间。"),
    ("averageSpeed", "拥堵路段平均速度。"), ("forecastLevel", "目标日期时段预测。"),
    ("forecastReason", "样本选择依据。"), ("noiseNote", "通勤、节假日或异常噪声说明。"),
    ("available", "是否具有可展示数据。"),
], [2500, 6860])
add_h2("13.4 HTTP 状态")
add_table(["状态", "场景"], [
    ("200", "查询或生成成功；外部部分能力不可用时可返回带说明的降级数据。"),
    ("204", "历史记录删除成功。"),
    ("400", "目的地、天数、预算或自然语言硬约束不可行。"),
    ("500", "未被业务异常捕获的服务端错误。"),
], [1600, 7760])

add_h1("14. 异常处理与降级策略", page_break=True)
add_table(["依赖", "异常", "降级策略"], [
    ("高德地理编码", "无结果或超时", "拒绝无效目的地；不使用虚构城市继续规划。"),
    ("高德 POI", "城市无景点", "尝试动态同步；仍为空则返回明确错误。"),
    ("高德路线", "路径失败", "使用估算转场并在地图显示虚线降级。"),
    ("DeepSeek", "Key 缺失或响应失败", "保留规则引擎生成的可执行行程。"),
    ("博查/网页", "搜索或抓取失败", "研究报告显示数据不足，不伪造证据。"),
    ("天气", "主源失败", "切换备用天气源；均失败时显示不可用。"),
    ("百度路况", "空结构、限流、无道路", "扩大半径、道路回退、保留最近成功值或明确无覆盖。"),
    ("数据库", "写入失败", "规划主响应不应因非关键研究快照失败而丢失；关键历史保存失败需记录错误。"),
], [1800, 2700, 4860], 8.6)
add_h2("14.1 日志要求")
add_bullets([
    "记录外部服务名称、业务对象、响应状态和降级路径，不记录 API Key。",
    "交通采样记录景点、参考道路、是否保存成功和限流情况。",
    "生产环境为每次规划生成 requestId，串联网搜、地图、AI 和持久化日志。",
])

add_h1("15. 安全与合规设计", page_break=True)
add_h2("15.1 密钥管理")
add_bullets([
    "后端密钥通过 AMAP_KEY、DEEPSEEK_API_KEY、BOCHA_API_KEY、BAIDU_TRAFFIC_AK 等环境变量注入。",
    "前端高德 JS Key 必须配置域名白名单和安全密钥；不得把后端 Web Service Key 发给浏览器。",
    "已在聊天、截图或日志中暴露的密钥应在平台控制台轮换。",
    "仓库、文档、数据库快照和错误响应不得保存真实密钥。",
])
add_h2("15.2 网络采集合规")
add_bullets([
    "优先使用官方开放 API、景区官网、政府/文旅公开页面和授权数据源。",
    "遵守 robots.txt、平台服务条款、著作权和个人信息保护要求。",
    "不绕过登录、验证码、设备指纹、签名、限流和访问控制。",
    "证据只保存必要摘要和 URL，不批量复制受版权保护的全文。",
])
add_h2("15.3 输入与输出安全")
add_p("对城市、日期、预算和天数执行服务端校验；自由文本只作为规划约束输入，不允许控制系统提示或外部请求地址。外部 URL 在抓取前进行协议、域名/IP 和重定向安全检查，防止 SSRF。")

add_h1("16. 部署与运行", page_break=True)
add_h2("16.1 环境要求")
add_table(["组件", "要求"], [
    ("JDK", "Java 21"), ("Maven", "3.9+ 推荐"), ("Node.js", "满足当前 Vite 版本"),
    ("MySQL", "8.x，utf8mb4"), ("浏览器", "现代 Chromium/Edge"),
], [2200, 7160])
add_h2("16.2 后端环境变量")
add_table(["变量", "用途", "必需性"], [
    ("SPRING_PROFILES_ACTIVE=mysql", "启用 MySQL profile", "生产必需"),
    ("DB_URL/DB_USERNAME/DB_PASSWORD", "数据库连接", "MySQL 必需"),
    ("AMAP_KEY", "高德 Web Service", "核心"),
    ("DEEPSEEK_API_KEY", "AI 解析与个性化", "可降级"),
    ("BOCHA_API_KEY", "正式网页搜索", "联网研究必需"),
    ("BAIDU_TRAFFIC_AK", "实时路况", "交通模块必需"),
    ("OPENWEATHER_API_KEY", "天气主源", "天气模块"),
    ("WEATHERAPI_KEY", "天气备用源", "建议"),
], [3000, 3800, 2560], 8.7)
add_h2("16.3 启动命令")
add_code("# 后端\ncd backend\nmvn spring-boot:run\n\n# 前端\ncd frontend\nnpm.cmd install\nnpm.cmd run dev")
add_h2("16.4 运行检查")
add_numbers([
    "访问 http://localhost:8080/api/health，确认返回 status=UP。",
    "访问 http://localhost:5173，确认首页可加载。",
    "提交真实城市，确认目的地校验、研究报告、三套行程和地图显示。",
    "五分钟后检查 traffic_snapshot 是否新增记录。",
])

add_h1("17. 测试设计", page_break=True)
add_h2("17.1 功能测试用例")
add_table(["编号", "场景", "输入/操作", "预期结果"], [
    ("TC-01", "有效目的地", "杭州，3 天", "生成三套方案。"),
    ("TC-02", "无效目的地", "不存在的字符串", "返回 400 并提示检查名称。"),
    ("TC-03", "指定地点", "第二天华山全天", "华山只进入第二天并占用全天。"),
    ("TC-04", "不合理日出", "下午 6 点爬山看日出", "拒绝并解释日出时间矛盾。"),
    ("TC-05", "登山恢复", "第二天全天登山", "第三天排除高强度登山。"),
    ("TC-06", "预算差异", "预算 3000", "目标为 2100/3000/3900，组合有明显差异。"),
    ("TC-07", "地铁换乘", "公共交通路线", "显示线路和上下车站。"),
    ("TC-08", "步行街路况", "河坊街", "道路名失败后返回周边区域路况。"),
    ("TC-09", "大景区路况", "西湖", "800m 空结果后扩大半径或回退主干道。"),
    ("TC-10", "历史恢复", "打开历史记录", "计划及研究快照与生成时一致。"),
], [900, 1600, 3000, 3860], 8.4)
add_h2("17.2 交通模块专项测试")
add_bullets([
    "对同一景点并发请求三次，确认只触发一次采样。",
    "模拟 status=302，确认不覆盖最近成功快照。",
    "模拟 status=0 但 evaluation、description、road_traffic 均为空，确认进入下一查询层。",
    "验证 07:00-10:00 和 17:00-20:00 的 peakHourNoise。",
    "验证近 60 天预测排除 status=0 和噪声样本。",
])
add_h2("17.3 验收标准")
add_table(["维度", "标准"], [
    ("正确性", "硬约束无遗漏；日期、时间、线路和地图按天一致。"),
    ("时效性", "研究证据在目标日前三个月；路况显示分钟级时间戳。"),
    ("可解释性", "热度有具体 URL；客流和交通预测有模型依据。"),
    ("稳定性", "单一外部 API 失败不导致整个页面不可用。"),
    ("持久性", "MySQL profile 下重启后历史、研究和交通快照仍存在。"),
], [2100, 7260])

add_h1("18. 性能与扩展设计", page_break=True)
add_h2("18.1 性能优化")
add_bullets([
    "高德步行、酒店和餐饮结果使用坐标与档位缓存。",
    "联网研究限制查询结果、抓取页数、单域数量和访问间隔。",
    "交通模块按 attractionId 锁定、请求限速、快照复用。",
    "数据库按城市、时间和景点建立索引，JSON 快照避免复杂级联写入。",
])
add_h2("18.2 后续扩展")
add_table(["方向", "建议"], [
    ("节假日识别", "接入国务院节假日数据，不仅以周末近似。"),
    ("客流模型", "积累 MySQL 历史样本后训练分城市、景点类型的时序模型。"),
    ("价格", "接入合法票务、酒店和团购开放 API，替换预算估算。"),
    ("账号体系", "增加用户表和鉴权，将历史记录绑定用户。"),
    ("异步任务", "将联网研究与批量交通采样迁移到消息队列和任务状态接口。"),
    ("可观测性", "增加 Micrometer、结构化日志、外部 API 成功率和配额告警。"),
], [2200, 7160])
add_h2("18.3 已知限制")
add_bullets([
    "百度路况覆盖机动车道路，偏远景区或纯步行区域可能没有可用数据。",
    "当前交通预测在样本不足时退化为最新状态，长期预测需持久 MySQL 数据。",
    "平台搜索结果和第三方价格会变化，出发前仍需在官方页面复核。",
    "高德 POI 评分不等同于携程、美团评分，跨平台深链需要对应平台开放 ID。",
])

add_h1("附录 A：关键配置基线", page_break=True)
add_table(["配置", "默认/建议值"], [
    ("服务端口", "8080"), ("前端端口", "5173"),
    ("旅行天数", "1-10 天"), ("预算范围", "500-200000 元"),
    ("研究窗口", "目标日前 3 个月"), ("最小证据量", "5 条"),
    ("交通刷新", "每 5 分钟"), ("交通半径", "800m -> 1000m -> 道路查询"),
    ("交通预测历史", "近 60 天同星期相邻小时"), ("酒店入住/退房", "13:00 后入住，次日 13:00 前退房"),
], [3000, 6360])
add_h1("附录 B：数据真实性声明")
add_p("系统提供的是辅助规划结果。景点开放、预约、限流、票价、交通、天气和商户营业状态可能变化，用户出发前应以景区官网、地图实时导航、气象服务及交易平台最终页面为准。系统界面应持续展示数据来源、采集时间和预测边界。")

# Prevent orphaned headings and set metadata.
doc.core_properties.title = "智能旅游助手详细设计文档"
doc.core_properties.subject = "基于通用智能体的全国智能旅游助手详细设计"
doc.core_properties.author = "智能旅游助手项目组"
doc.core_properties.keywords = "Vue3, Spring Boot, MySQL, 高德地图, 百度路况, DeepSeek, 旅游规划"
doc.core_properties.comments = "由当前项目代码基线生成，不包含任何真实 API Key。"

doc.save(OUT)
print(OUT)
