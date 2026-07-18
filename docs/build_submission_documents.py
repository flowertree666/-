from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_internship_report import (
    GREEN, MUTED, NAVY, add_bullet, add_header_footer, add_heading, add_note,
    add_number, add_run_paragraph, add_table, set_font, style_paragraph,
)


OUT_DIR = Path(__file__).resolve().parent / "submission_materials"
OUT_DIR.mkdir(parents=True, exist_ok=True)

TEAM = [
    ("组长", "程鹏", "32301108", "计算机2302", "总体方案、南京数据集、RAG 接入、后端集成、关键算法与最终联调。"),
    ("成员 A", "马骁泽", "32301115", "计算机2302", "北京旅游 POI、官方证据与知识片段整理，完成数据规范化与质量检查。"),
    ("成员 B", "柴正宇", "32301106", "计算机2302", "杭州旅游 POI、官方证据与知识片段整理，补充季节性、预约和开放规则。"),
    ("成员 C", "梅佳俊", "32301116", "计算机2302", "Vue 3 交互页面、行程时间线、地图可视化和多方案展示优化。"),
    ("成员 D", "杨宇韬", "32301124", "计算机2302", "测试用例、异常场景验证、作品演示素材、PPT 与答辩材料整理。"),
]


def new_doc(running_title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    add_header_footer(section)
    section.header.paragraphs[0].clear()
    set_font(section.header.paragraphs[0].add_run(running_title), 8.5, False, MUTED)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    return doc


def cover(doc, title, subtitle, document_type):
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    style_paragraph(p, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("基于通用智能体的智能旅游助手"), 15, True, GREEN)
    p = doc.add_paragraph()
    style_paragraph(p, after=14, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run(title), 26, True, NAVY)
    p = doc.add_paragraph()
    style_paragraph(p, after=30, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run(subtitle), 12.5, False, MUTED)
    add_table(doc, ["材料类型", document_type], [
        ("项目名称", "智游 Agent"),
        ("参赛方向", "基于大模型、RAG 与地图服务的个性化旅行规划"),
        ("项目组长", "程鹏 / 32301108 / 计算机2302"),
        ("团队规模", "5 人"),
        ("提交时间", "2026 年 7 月"),
    ], [2200, 7160])
    add_note(doc, "核心主张", "以可追溯知识为依据，以明确约束为底线，生成真正可执行的个性化旅行路线。")


def save(doc, filename, title, subject):
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "智游 Agent 项目组"
    doc.core_properties.keywords = "智能旅游,RAG,Vue3,Spring Boot,高德地图,DeepSeek"
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_product_intro():
    doc = new_doc("智游 Agent | 作品简介")
    cover(doc, "作品简介", "从用户一句话需求到可落地的多日行程", "比赛作品说明")

    add_heading(doc, "一、项目背景", 1, page_break=True)
    add_run_paragraph(doc, "自由行用户常常需要在景区官网、地图、攻略社区和票务平台之间反复切换。景点信息分散、开放规则变化、路线安排缺少空间与体力约束，导致攻略制作耗时且最终方案不一定可执行。尤其当用户提出“第一天去八达岭长城”“带老人同行”“早起看日出”等要求时，普通模板攻略难以稳定落实。")
    add_heading(doc, "二、作品定位", 1)
    add_run_paragraph(doc, "智游 Agent 是面向国内自由行场景的智能旅游助手。系统以城市 RAG 知识库为可信信息底座，以规则约束与地图路线为执行保障，再由 DeepSeek 完成自然语言理解和个性化表达，形成经济、标准、舒适三套差异化路线。")
    add_note(doc, "目标用户", "希望减少攻略时间、重视路线合理性、需要照顾老人或儿童、希望控制预算的自由行用户。")

    add_heading(doc, "三、核心功能", 1)
    add_table(doc, ["功能", "用户价值", "实现方式"], [
        ("需求画像", "输入城市、天数、日期、预算、偏好与特别期待。", "前端表单校验 + DeepSeek / 本地规则解析。"),
        ("RAG 检索", "获得有来源、有有效期的景点规则与建议。", "三地 JSONL 数据 + ChromaDB 检索。"),
        ("多方案规划", "比较经济、标准、舒适三套路线。", "预算分档 + 地理聚类 + 每日时间约束。"),
        ("真实路线", "查看按日分离的地图、步行、公交/地铁或打车路线。", "高德地图 Web 服务与 JS API。"),
        ("完整时间线", "起床、三餐、转场、景点、住宿与行李安排连续呈现。", "动态时间游标 + 餐饮住宿 POI。"),
        ("历史记录", "保存并重新查看生成结果和规划依据。", "MySQL / H2 持久化。"),
    ], [1500, 3900, 3960])

    add_heading(doc, "四、典型使用流程", 1, page_break=True)
    add_number(doc, "用户输入目的地、旅行天数、出发日期、预算和偏好。")
    add_number(doc, "系统校验目的地；若本地数据集覆盖该城市，则优先使用数据集 POI，否则回退到高德搜索。")
    add_number(doc, "RAG 服务检索开放、预约、体力、季节和数据治理规则，并返回来源与有效期。")
    add_number(doc, "路线引擎落实必去点、指定日期、每日时间、体力恢复、空间距离和预算约束。")
    add_number(doc, "高德服务补充交通路线、餐饮和住宿，DeepSeek生成个性化建议，页面按日展示完整方案。")
    add_heading(doc, "五、作品特色", 1)
    add_bullet(doc, "数据与路线一致：RAG 检索和路线选点使用同一套城市 POI，减少“介绍了却没安排”的问题。")
    add_bullet(doc, "明确需求是硬约束：指定景点和第几天优先执行，模型不可用时仍由本地规则保障。")
    add_bullet(doc, "证据可追溯：知识片段显示来源 URL、有效期、事实类型和相关性。")
    add_bullet(doc, "结果可执行：时间线覆盖吃、住、行、游与行李处理，不只输出景点清单。")
    add_bullet(doc, "服务可降级：地图、天气、RAG 或模型异常时保留基础规划并明确提示。")

    add_heading(doc, "六、当前成果与边界", 1)
    add_table(doc, ["成果", "现状"], [
        ("城市知识库", "覆盖南京、北京、杭州，共 298 条可信知识片段。"),
        ("景点数据", "三地标准化 POI 同步进入路线数据库并参与选点。"),
        ("核心链路", "需求输入、RAG 检索、路线生成、地图、时间线、历史记录已贯通。"),
        ("实验交付", "提供含 Java 运行时、H2 演示库和向量索引的一键启动包。"),
        ("数据边界", "票价、开放、限流与实时客流仍应在出发前以景区官方信息复核。"),
    ], [2200, 7160])
    return save(doc, "01-作品简介-智游Agent.docx", "智游 Agent 作品简介", "比赛作品说明")


def build_team_intro():
    doc = new_doc("智游 Agent | 团队介绍")
    cover(doc, "团队介绍", "五人协作完成数据、算法、交互与答辩闭环", "团队能力与分工说明")

    add_heading(doc, "一、团队概况", 1, page_break=True)
    add_run_paragraph(doc, "团队由计算机 2302 班 5 名成员组成，围绕“数据建设、系统研发、体验设计、质量验证”四条工作线协同推进。组长负责统一架构、接口、数据标准与最终验收，成员按城市数据、前端交互和测试材料形成明确交付物。")
    add_table(doc, ["角色", "姓名", "学号", "班级", "主要工作与交付物"], [
        tuple(x for x in member) for member in TEAM
    ], [800, 1100, 1300, 1300, 4860])

    add_heading(doc, "二、成员工作说明", 1)
    for role, name, sid, clazz, work in TEAM:
        add_heading(doc, f"{role}：{name}", 2)
        add_run_paragraph(doc, work, first_indent=False)

    add_heading(doc, "三、协作机制", 1, page_break=True)
    add_heading(doc, "3.1 统一数据契约", 2)
    add_run_paragraph(doc, "三地数据均采用 pois.jsonl、knowledge_chunks.jsonl、sources.jsonl 结构，统一 POI 编号、城市、坐标、标签、事实类型、来源 URL 和有效期。各成员先完成本地自查，再由组长执行导入、检索和路线命中测试。")
    add_heading(doc, "3.2 接口并行开发", 2)
    add_run_paragraph(doc, "前端围绕稳定的计划请求和响应结构开发，后端负责路线、预算、RAG 和地图接口。接口字段变化由组长统一确认，避免成员之间反复修改。")
    add_heading(doc, "3.3 以场景验收", 2)
    add_run_paragraph(doc, "团队采用真实语句作为验收样例，例如“第一天去八达岭长城”“带老人避免高强度步行”。只有当数据命中、路线安排和页面展示同时正确时，功能才视为完成。")
    add_heading(doc, "3.4 材料同步", 2)
    add_run_paragraph(doc, "测试负责人同步沉淀截图、测试结果和演示步骤，避免答辩材料在项目完成后重新整理。团队文档统一使用同一套架构、数据规模和功能表述。")

    add_heading(doc, "四、团队贡献矩阵", 1)
    add_table(doc, ["工作域", "牵头成员", "协作成员", "验收成果"], [
        ("南京数据与 RAG", "程鹏", "杨宇韬", "南京 POI、知识片段、向量索引与检索测试。"),
        ("北京数据", "马骁泽", "程鹏", "北京 POI、官方证据、字段规范化与路线命中。"),
        ("杭州数据", "柴正宇", "程鹏", "杭州 POI、季节/预约规则与可信来源整理。"),
        ("前端体验", "梅佳俊", "程鹏", "需求表单、地图、时间线、方案切换与历史页面。"),
        ("测试与答辩", "杨宇韬", "全体成员", "测试用例、演示视频、PPT 和提交材料。"),
    ], [1800, 1500, 1500, 4560])
    add_note(doc, "团队特点", "分工不是孤立模块：城市数据必须进入同一 RAG 与路线链路，前端展示必须通过真实后端结果验收，测试材料必须能复现核心场景。")
    return save(doc, "02-团队介绍-智游Agent.docx", "智游 Agent 团队介绍", "团队能力与分工说明")


def build_technical_solution():
    doc = new_doc("智游 Agent | 技术方案文档")
    cover(doc, "技术方案文档", "可信知识检索、约束路线规划与地图能力融合", "系统技术方案")

    add_heading(doc, "一、建设目标与设计原则", 1, page_break=True)
    add_run_paragraph(doc, "项目目标是把用户的自然语言旅行需求转化为可解释、可执行、可复核的多日行程。系统必须同时满足个性化、时间与空间合理性、数据可追溯、外部服务可降级和结果可展示五项要求。")
    add_bullet(doc, "规则兜底：明确地点、日期和体力约束由确定性逻辑保障。")
    add_bullet(doc, "数据同源：检索知识与路线候选来自同一套城市 POI。")
    add_bullet(doc, "动态校正：交通、天气、餐饮和住宿由地图与天气服务补充。")
    add_bullet(doc, "证据治理：事实片段保留来源、类型、有效期和可信状态。")
    add_bullet(doc, "分层降级：外部 API 不可用时仍能生成基础路线。")

    add_heading(doc, "二、总体架构", 1)
    add_table(doc, ["层次", "技术组件", "主要职责"], [
        ("交互层", "Vue 3 + Vite", "采集需求，展示多方案、RAG 证据、地图、时间线和历史记录。"),
        ("业务层", "Spring Boot / Java 21", "目的地校验、意图解析、路线规划、预算、历史与接口编排。"),
        ("知识层", "FastAPI + ChromaDB", "加载三地知识片段，执行城市过滤、向量检索与有效期治理。"),
        ("数据层", "MySQL / H2 + JSONL", "保存 POI、来源、历史计划和标准化城市数据。"),
        ("能力层", "高德、百度、DeepSeek、天气 API", "地图 POI、真实路线、路况、语言理解与天气信息。"),
    ], [1300, 2700, 5360])

    add_heading(doc, "三、数据集与 RAG 设计", 1, page_break=True)
    add_heading(doc, "3.1 数据结构", 2)
    add_table(doc, ["文件", "内容", "关键字段"], [
        ("pois.jsonl", "可参与路线规划的景点实体。", "poi_id、name、city、坐标、标签、体力等级、建议时长。"),
        ("knowledge_chunks.jsonl", "用于检索的事实和规则片段。", "chunk_id、content、fact_type、source_url、valid_until。"),
        ("sources.jsonl", "来源登记与可信状态。", "source_id、URL、发布主体、发布日期、验证状态。"),
    ], [1800, 2700, 4860])
    add_heading(doc, "3.2 索引与检索", 2)
    add_number(doc, "启动 RAG 服务时扫描包含 POI 与知识片段的城市目录并加载 ChromaDB。")
    add_number(doc, "将标题、正文、标签和城市信息共同编码，检索时按目标城市过滤。")
    add_number(doc, "对命中结果进行名称别名加权、事实类型排序和有效期过滤。")
    add_number(doc, "Java 后端将知识片段作为规划上下文，并把来源、得分与有效期返回前端。")
    add_note(doc, "数据规模", "当前覆盖南京、北京、杭州 3 个城市，共 298 条可信知识片段；不可信或仅占位的片段不会进入有效检索。")

    add_heading(doc, "四、路线规划引擎", 1, page_break=True)
    add_heading(doc, "4.1 输入约束", 2)
    add_table(doc, ["约束", "处理方式"], [
        ("旅行天数", "严格生成对应天数，不允许缺天或多天。"),
        ("必去景点", "从模型抽取与本地 POI 名称/别名匹配中合并为硬约束。"),
        ("指定日期", "将景点绑定到用户指定的第几天，超出行程天数时报错。"),
        ("体力恢复", "登山或全天高强度活动后降低次日高强度项目密度。"),
        ("空间距离", "按片区聚类，计算转场距离并减少跨区折返。"),
        ("时间窗口", "动态推进起床、早餐、交通、游玩、午餐、晚餐和住宿时间。"),
        ("预算", "按经济、标准、舒适档位调整住宿、餐饮、交通和组合。"),
    ], [2200, 7160])
    add_heading(doc, "4.2 生成策略", 2)
    add_run_paragraph(doc, "路线引擎首先选取每日锚点，再从同片区候选中按偏好、热度、季节适配、距离和方案差异评分补充景点。已绑定日期的必去点优先成为当天锚点；候选不足或外部路线失败时，系统保留可解释的估算与缓冲时间。")

    add_heading(doc, "五、地图、交通、餐饮与住宿", 1, page_break=True)
    add_heading(doc, "5.1 高德地图能力", 2)
    add_bullet(doc, "目的地校验与城市标准化，阻止不存在或错误目的地继续规划。")
    add_bullet(doc, "POI 搜索与景点评分，用于城市无数据集时的全国回退。")
    add_bullet(doc, "步行、驾车和公共交通路线；公交结果解析地铁线路、上下车站和换乘段。")
    add_bullet(doc, "餐饮与酒店周边搜索，为早餐、午餐、晚餐和住宿提供候选。")
    add_bullet(doc, "前端 JS API 按天绘制地图，切换日期时只展示当前日路线。")
    add_heading(doc, "5.2 住宿与行李规则", 2)
    add_run_paragraph(doc, "景点集中时优先保持同一住宿点；跨片区距离较大时允许换酒店。系统按 13:00 入住、次日 13:00 前退房的规则生成时间线，并在退房后安排酒店寄存或携带行李转移。多个住宿点之间的交通也进入路线与预算。")
    add_heading(doc, "5.3 路况和客流参考", 2)
    add_run_paragraph(doc, "百度实时路况用于描述景点周边道路压力并保存采样时间。系统标记早晚高峰、周末和节假日等噪音项，不把道路拥堵直接等同于景区人数，仅作为预计拥挤度的辅助变量。")

    add_heading(doc, "六、大模型与个性化生成", 1, page_break=True)
    add_run_paragraph(doc, "DeepSeek 承担两类任务：一是把自由文本提取为必去点、排除点、指定日期、全天、日出、登山和起床时间等结构化约束；二是在完成规则校验的路线草案上生成每日主题和个性化建议。")
    add_note(doc, "稳定性设计", "模型输出不是路线约束的唯一来源。本地 POI 目录会再次识别明确景点及其附近的“第几天”表达，确保模型未配置或抽取失败时仍能执行关键需求。")
    add_heading(doc, "七、前端交互设计", 1)
    add_table(doc, ["页面区域", "交互设计"], [
        ("需求表单", "目的地、天数、日期、预算、偏好和自由文本；无效目的地即时提示。"),
        ("方案区", "经济、标准、舒适三档切换并显示预算区间。"),
        ("证据区", "折叠展示 RAG 片段、来源链接、类型、得分与有效期。"),
        ("地图区", "按天切换真实道路与景点标记，支持缩放。"),
        ("时间线", "连续展示起床、餐饮、交通、景点、住宿和费用。"),
        ("历史记录", "保存计划输入、生成结果和当时的规划证据。"),
    ], [2200, 7160])

    add_heading(doc, "八、数据库与接口", 1, page_break=True)
    add_table(doc, ["数据对象", "主要内容", "用途"], [
        ("attraction", "名称、城市、坐标、标签、热度、来源、开放与票价。", "路线候选与景点展示。"),
        ("travel_history", "用户输入、计划 JSON、研究快照、创建时间。", "历史记录与复现。"),
        ("traffic_snapshot", "道路状态、采样时间、噪音标签。", "周边交通和趋势参考。"),
        ("research_insight", "城市近期研究报告和证据列表。", "联网研究兼容与历史留存。"),
    ], [1900, 4800, 2660])
    add_heading(doc, "8.1 核心接口", 2)
    add_bullet(doc, "POST /api/plans：生成三套个性化旅行计划。")
    add_bullet(doc, "GET /api/rag/search：按城市、日期和查询词检索知识片段。")
    add_bullet(doc, "GET /api/amap/validate：校验目的地并返回标准城市名。")
    add_bullet(doc, "GET /api/history：查询历史计划；DELETE /api/history/{id}：删除记录。")

    add_heading(doc, "九、异常处理、安全与合规", 1, page_break=True)
    add_table(doc, ["风险", "控制措施"], [
        ("无效目的地", "高德地理编码与本地数据集共同校验，校验失败直接报错。"),
        ("模型不可用", "规则引擎和本地 POI 硬约束继续生成路线。"),
        ("RAG 不可用", "路线数据库继续工作，前端显示知识服务降级。"),
        ("地图调用失败", "使用距离和时间估算，并在地图中以降级状态提示。"),
        ("数据过期", "动态公告设置有效期，过期片段降权或排除。"),
        ("平台数据合规", "不绕过抖音、小红书等平台反爬；优先使用官网、开放 API 和授权数据。"),
        ("密钥安全", "正式部署应使用环境变量、域名白名单和后端代理；实验密钥不得公开发布。"),
    ], [2200, 7160])

    add_heading(doc, "十、测试与验收", 1)
    add_table(doc, ["测试类型", "典型用例", "验收标准"], [
        ("功能", "北京 3 日、第一天去八达岭长城。", "第 1 天必须包含八达岭长城。"),
        ("RAG", "查询中山陵预约与开放规则。", "Top 结果命中对应 POI，并返回来源和有效期。"),
        ("路线", "多日公共交通方案。", "地图按日分离，时间线包含地铁线路与转场。"),
        ("异常", "输入不存在的城市。", "停止规划并提示目的地错误。"),
        ("预算", "输入 3000 元。", "三档目标约为 70%、100%、130%，组合存在差异。"),
        ("历史", "生成后重新打开记录。", "完整恢复方案、地图数据和规划证据。"),
    ], [1500, 4300, 3560])

    add_heading(doc, "十一、部署与演示", 1, page_break=True)
    add_run_paragraph(doc, "项目提供前后端源代码与 18-5-test 实验包。实验包内置 Java 21、Spring Boot JAR、Vue 静态资源、H2 数据库、三地数据集和 Chroma 向量索引。组员解压后运行启动脚本即可打开系统；Python 环境可用于启动完整 RAG 服务。")
    add_heading(doc, "十二、后续扩展", 1)
    add_bullet(doc, "扩大到更多城市，并建立自动化来源更新、过期检查和人工复核流程。")
    add_bullet(doc, "从字符 n-gram 向量升级到中文语义嵌入模型，并完善离线评测集。")
    add_bullet(doc, "接入授权酒店、票务与预约数据，提高价格和开放状态的实时性。")
    add_bullet(doc, "细化老人、儿童、行李和体力恢复模型，提升行程舒适度。")
    add_bullet(doc, "将系统部署到云端，使用容器编排、密钥管理、缓存与调用监控。")
    add_note(doc, "结论", "该方案以数据治理和约束规划为稳定底座，以 RAG 和大模型增强个性化，以地图服务保障可执行性，适合课程答辩与后续工程扩展。")
    return save(doc, "03-技术方案文档-智游Agent.docx", "智游 Agent 技术方案文档", "系统技术方案")


def build_demo_script():
    doc = new_doc("智游 Agent | 功能演示视频录制脚本")
    cover(doc, "功能演示视频录制脚本", "建议时长 4 分钟，按真实项目操作完成录屏", "演示视频制作辅助材料")
    add_heading(doc, "一、录制要求", 1, page_break=True)
    add_bullet(doc, "建议分辨率 1920×1080，浏览器缩放 100%，关闭无关窗口和通知。")
    add_bullet(doc, "录制前启动后端与 RAG 服务，确认顶部显示“知识服务已连接”。")
    add_bullet(doc, "使用北京 3 日样例：预算 6000 元，特别期待填写“第一天去八达岭长城”。")
    add_bullet(doc, "讲解时强调数据来源、硬约束与降级边界，不声称客流和票价绝对实时。")
    add_heading(doc, "二、分镜与讲解词", 1)
    add_table(doc, ["时间", "画面操作", "建议讲解"], [
        ("00:00-00:20", "显示项目名称和首页。", "智游 Agent 是融合 RAG、大模型和高德地图的个性化旅行规划系统。"),
        ("00:20-00:50", "展示知识库 298 条、3 城覆盖和规划方法。", "系统优先检索本地可信知识，再结合路线、天气、住宿与交通安排。"),
        ("00:50-01:20", "输入北京、3 天、6000 元和“第一天去八达岭长城”。", "特别期待会被解析成地点和日期硬约束，而不是只作为文案参考。"),
        ("01:20-01:50", "生成后展开 RAG 规划依据。", "每条依据保留事实类型、POI 编号、来源和复核入口。"),
        ("01:50-02:20", "切换经济、标准、舒适三套方案。", "三档方案会调整住宿、餐饮、交通、景点组合和节奏。"),
        ("02:20-02:50", "切换第 1/2/3 天地图。", "地图按天分离并优先绘制高德真实路线，失败时明确显示降级状态。"),
        ("02:50-03:30", "滚动到第 1 天时间线和八达岭长城。", "时间线从起床开始，连续安排三餐、交通、休息和游玩；第一天约束已被明确兑现。"),
        ("03:30-03:50", "展示历史记录和备选景点。", "历史记录可恢复计划与依据，备选景点用于下次重新生成参考。"),
        ("03:50-04:00", "回到项目标题或结束页。", "系统的核心不是随机拼接，而是生成有依据、守约束、可执行的路线。"),
    ], [1250, 3150, 4960])
    add_heading(doc, "三、提交前检查", 1, page_break=True)
    add_bullet(doc, "视频内没有显示 API Key、数据库密码或个人隐私信息。")
    add_bullet(doc, "八达岭长城确实出现在第 1 天时间线，而不是只出现在 AI 建议中。")
    add_bullet(doc, "RAG 证据、三档方案、按日地图、时间线和历史记录均有画面。")
    add_bullet(doc, "字幕和讲解中的数据规模统一为：3 个城市、75 个标准化 POI、298 条可信片段。")
    add_bullet(doc, "最终文件建议命名：04-功能演示视频-智游Agent.mp4。")
    add_note(doc, "注意", "本脚本是录制辅助材料，不能替代比赛要求的 MP4 功能演示视频。")
    return save(doc, "04-功能演示视频录制脚本-智游Agent.docx", "智游 Agent 功能演示视频录制脚本", "视频分镜与讲解词")


if __name__ == "__main__":
    for output in (build_product_intro(), build_team_intro(), build_technical_solution(), build_demo_script()):
        print(output)
