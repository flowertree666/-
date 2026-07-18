from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "实习总结报告-组长-智能旅游助手.docx"

NAVY = "173D30"
GREEN = "2F6F57"
ORANGE = "E47D3D"
MUTED = "66736D"
LIGHT = "EFF5F1"
TABLE_FILL = "E8F0EC"


def set_font(run, size=10.5, bold=None, color="1F2933", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = table_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    table_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_paragraph(paragraph, before=0, after=6, line=1.18, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_run_paragraph(doc, text, *, size=10.5, bold=False, color="1F2933", before=0, after=6,
                      line=1.25, align=None, first_indent=True):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, line, align)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 1:
        style_paragraph(p, before=16, after=8, line=1.1)
        r = p.add_run(text)
        set_font(r, 16, True, GREEN)
    else:
        style_paragraph(p, before=10, after=5, line=1.1)
        r = p.add_run(text)
        set_font(r, 12, True, NAVY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=4, line=1.18)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    r = p.add_run(text)
    set_font(r, 10.5, False)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, before=0, after=4, line=1.18)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    r = p.add_run(text)
    set_font(r, 10.5, False)
    return p


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    style_paragraph(p, before=1, after=1, line=1.18)
    r = p.add_run(title + "  ")
    set_font(r, 10.5, True, GREEN)
    r = p.add_run(text)
    set_font(r, 10.5, False, "365346")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        style_paragraph(p, before=0, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_font(p.add_run(header), 9.5, True, NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.1,
                            align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            set_font(p.add_run(value), 9.2, False)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, before=0, after=0, line=1.0)
    r = p.add_run("智能旅游助手项目 | 移动互联网短学期实习总结报告")
    set_font(r, 8.5, False, MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=0, after=0, line=1.0)
    r = p.add_run("课程实习材料")
    set_font(r, 8.5, False, MUTED)


def make_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    # Page 1: cover and basic registration fields.
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=12, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("移动互联网系统实习"), 15, True, GREEN)
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=18, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("实习总结报告"), 28, True, NAVY)
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=32, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("项目组长版"), 13, False, MUTED)
    add_table(doc, ["项目题目", "智游 Agent - 基于通用智能体的智能旅游助手"], [
        ("项目方向", "大模型与智能体赋能的个性化旅游路线规划"),
        ("承担角色", "项目组长 / 系统架构与集成负责人"),
        ("学号", "________________"),
        ("姓名", "________________"),
        ("班级", "________________"),
    ], [2300, 7060])
    add_note(doc, "提交提示", "请在打印或上传前补全学号、姓名、班级及第 2.3 节中的成员真实信息。")
    p = doc.add_paragraph()
    style_paragraph(p, before=60, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("2026 年 7 月"), 11, False, MUTED)

    # Page 2: purpose, group, product.
    add_heading(doc, "一、实习目的和要求", 1, page_break=True)
    add_run_paragraph(doc, "本次短学期实习以移动互联网应用开发为载体，要求学生在团队协作中完成从需求分析、数据处理、系统设计到部署演示的完整工程流程。项目聚焦旅游场景中信息分散、攻略同质化、路线不合理和动态信息难以追溯等问题，通过大模型、检索增强生成（RAG）和地图服务形成可解释的个性化出行方案。")
    add_run_paragraph(doc, "作为组长，我将实习目标具体落实为：建立可协作的任务边界；把用户自然语言需求转化为可执行约束；让景点知识、路线规划和展示结果共享同一套数据依据；并通过降级策略保证外部服务异常时系统仍可完成基础规划。")

    add_heading(doc, "二、项目组人员和具体分工", 1)
    add_heading(doc, "2.1 项目题目", 2)
    add_run_paragraph(doc, "智游 Agent - 基于通用智能体的智能旅游助手", first_indent=False)
    add_heading(doc, "2.2 所选题目", 2)
    add_run_paragraph(doc, "（十八）基于通用智能体的智能旅游助手开发", first_indent=False)
    add_heading(doc, "2.3 项目组人员和具体分工", 2)
    add_table(doc, ["成员", "姓名/学号/班级", "主要工作与交付物"], [
        ("组长", "请填写", "总体方案、南京数据集、RAG 接入、后端集成、关键算法与最终联调。"),
        ("成员 A", "请填写", "北京旅游 POI、官方证据与知识片段整理，完成数据规范化与质量检查。"),
        ("成员 B", "请填写", "杭州旅游 POI、官方证据与知识片段整理，补充季节性、预约和开放规则。"),
        ("成员 C", "请填写", "Vue 3 交互页面、行程时间线、地图可视化和多方案展示优化。"),
        ("成员 D", "请填写", "测试用例、异常场景验证、作品演示素材、PPT 与答辩材料整理。"),
    ], [900, 2200, 6260])
    add_note(doc, "组长协作机制", "采用“数据先行、接口并行、集成验收”的节奏：成员完成的数据和界面必须通过统一字段、接口和测试样例接入主分支，避免各自实现无法联调。")

    add_heading(doc, "三、作品简介", 1)
    add_run_paragraph(doc, "智游 Agent 面向国内自由行用户。用户输入目的地、旅行天数、出发日期、预算、偏好和特别期待后，系统优先检索对应城市的 RAG 知识库，并从同一套 POI 数据中选择景点，再结合高德地图路线、餐饮住宿、天气和交通信息，生成经济、标准、舒适三套可执行路线。页面以按日地图、动态时间线、费用估算、替代景点和 RAG 证据的形式展示结果。")

    # Page 3: technical solution and leader work.
    add_heading(doc, "四、技术方案文档", 1, page_break=True)
    add_heading(doc, "4.1 系统架构", 2)
    add_run_paragraph(doc, "系统采用前后端分离架构。前端使用 Vue 3 与 Vite，负责需求采集、地图和行程展示；后端使用 Java 21 与 Spring Boot，负责编排路线、调用外部能力、保存历史记录；生产环境使用 MySQL，演示环境可使用内置 H2；Python FastAPI 与 ChromaDB 组成轻量 RAG 服务。")
    add_table(doc, ["层次", "关键技术", "承担职责"], [
        ("前端", "Vue 3、Vite、高德 JS API", "表单校验、多方案切换、按日地图、时间线与证据折叠展示。"),
        ("后端", "Java 21、Spring Boot", "约束解析、路线生成、预算、历史记录、外部 API 适配。"),
        ("数据", "MySQL / H2、JSONL", "景点、来源、计划历史和三地标准化 POI 数据。"),
        ("RAG", "FastAPI、ChromaDB", "检索可信知识片段，过滤过期或不可信证据，返回来源链接。"),
        ("外部能力", "高德、百度、DeepSeek、天气服务", "POI、真实道路、路况、自然语言解析与天气辅助。"),
    ], [1000, 2600, 5760])
    add_heading(doc, "4.2 路线规划与约束处理", 2)
    add_number(doc, "目的地校验：先校验城市有效性；若数据集覆盖该城市，则优先使用数据集 POI，否则回退到高德 POI 搜索。")
    add_number(doc, "需求解析：将“第一天去八达岭长城”“第二天全天徒步”等自由文本转为必去景点、指定日期、全天游玩、体力强度和起床时间等硬约束。")
    add_number(doc, "路线生成：按预算构造经济、标准、舒适三种策略，在每日时间上限、景点地理距离、体力恢复和交通时间的共同约束下选点。")
    add_number(doc, "结果补全：补充高德步行、公交/地铁或驾车路线，安排餐饮、住宿、天气和费用估算；服务异常时明确显示降级信息，而不是伪造实时数据。")
    add_heading(doc, "4.3 本人作为组长的核心工作", 2)
    add_bullet(doc, "制定数据规范：统一 POI、知识片段、来源与有效期字段，完成南京 30 个核心景点的数据采集、标准化和可检索化。")
    add_bullet(doc, "完成 RAG 链路：建立 ChromaDB 本地向量索引，设计城市过滤、有效期过滤、来源可追溯和命中率测试流程。")
    add_bullet(doc, "负责 Java 后端：完成路线约束、三档预算、餐饮住宿、历史记录、高德路线与异常降级等核心逻辑的联调。")
    add_bullet(doc, "解决关键缺陷：针对“指定八达岭长城但未被安排”的问题，补充基于本地 POI 目录的必去点识别，使模型不可用时仍能稳定执行明确地点约束。")

    # Page 4: innovation and effect.
    add_heading(doc, "五、创新点", 1, page_break=True)
    add_heading(doc, "5.1 数据集、检索与路线共用同一事实来源", 2)
    add_run_paragraph(doc, "本项目不是仅把 RAG 用于生成介绍文案，而是将标准化 POI 同时同步到路线规划器。这样“检索到的南京、北京、杭州知识”和“参与排程的景点”保持一致，能够降低推荐景点与路线景点不一致的问题。")
    add_heading(doc, "5.2 自然语言需求转化为硬约束", 2)
    add_run_paragraph(doc, "对于明确提出的地点、日期和活动强度，系统优先保证约束满足，而不是只将其作为生成文案的参考。用户输入“第一天去八达岭长城”后，该景点会被固定在第一天；对于与日出时间矛盾等无法满足的要求，系统返回可解释提示。")
    add_heading(doc, "5.3 可解释与可降级的动态信息展示", 2)
    add_run_paragraph(doc, "RAG 结果展示来源 URL、有效期、事实类型和相关得分；地图、天气、路况、模型服务均设计了失败回退路径。系统将交通压力定位为辅助参考，不将道路拥堵直接等同于景区人数，避免把推测结果表述为实时事实。")
    add_heading(doc, "5.4 多目标预算方案", 2)
    add_run_paragraph(doc, "三种方案不只是对总价进行比例缩放，而是分别调整酒店、餐饮、交通与景点组合，形成经济、标准、舒适三组可比较的选择，并在费用不确定时标明估算依据。")

    add_heading(doc, "六、作品效果", 1)
    add_table(doc, ["验证场景", "预期效果", "结果"], [
        ("三地数据集检索", "南京、北京、杭州按城市过滤，返回带来源的知识片段。", "完成"),
        ("指定景点约束", "输入“第一天去八达岭长城”，第一天必须包含该点。", "完成"),
        ("多日路线展示", "地图按天分离，时间线仅显示当前日行程。", "完成"),
        ("预算差异", "经济、标准、舒适方案在住宿、餐饮和节奏上存在可见差异。", "完成"),
        ("服务降级", "RAG 或地图异常时保留基础规划并提示数据状态。", "完成"),
    ], [1600, 5700, 2060])
    add_note(doc, "演示建议", "答辩时依次展示：输入需求 -> RAG 规划依据 -> 三套路线 -> 按日地图和时间线 -> 历史记录。重点说明“知识检索、路线选择、约束执行”使用同一套 POI 数据。")

    # Page 5: reflection.
    add_heading(doc, "七、实习总结和心得", 1, page_break=True)
    add_run_paragraph(doc, "通过本次实习，我对“能运行的功能”和“可交付的工程系统”之间的差别有了更具体的认识。旅游规划不是简单地把热门景点串联起来，而是需要同时处理用户明确要求、空间距离、时间窗口、体力消耗、预算和数据时效性。作为组长，我首先承担的是把模糊目标拆成可验收任务，再让数据、算法和界面围绕同一份约束工作。")
    add_run_paragraph(doc, "在技术实现中，我最深的体会是要把大模型放在合适的位置。大模型适合解析自然语言、补充解释和生成个性化建议，但不应该替代明确的业务约束。因此，我在系统中保留了确定性的目的地校验、POI 匹配、每日排程和异常处理。当模型未配置或响应失败时，用户提出的“指定地点、指定日期”等关键要求依然能够执行。这种“模型增强而非模型依赖”的设计提高了系统的稳定性。")
    add_run_paragraph(doc, "RAG 数据集建设也让我认识到数据质量比数据数量更重要。我们为景点记录来源、有效期、事实类型和验证说明，并对过期或不可信片段降权或排除。这样的设计既能支持检索，也能在答辩中解释推荐依据，避免把旧公告、模糊评价或推测性结论误当作实时事实。")
    add_run_paragraph(doc, "在团队协作方面，我负责确定统一数据格式、接口边界和验收用例，成员分别完成北京、杭州数据、前端交互与测试材料。实践证明，早期先约定字段和交付标准可以显著降低后期合并成本；组长不能只完成自己的代码，还要及时识别接口不一致、数据重复和功能遗漏，并推动形成可演示的闭环。")
    add_run_paragraph(doc, "后续我计划继续扩展更多城市数据，使用更成熟的向量检索和评测集，并将酒店价格、景区预约和交通预测改为更可验证的授权数据源。同时会进一步细化用户画像和体力恢复模型，使路线在满足偏好的同时更贴近真实旅行体验。本次实习提升了我在全栈开发、数据治理、系统集成和工程沟通方面的能力，也让我更加重视隐私、合规和结果可解释性。")

    doc.core_properties.title = "实习总结报告-组长-智能旅游助手"
    doc.core_properties.subject = "移动互联网系统短学期实习总结"
    doc.core_properties.author = "项目组长"
    doc.core_properties.keywords = "智能旅游助手,RAG,Vue3,Spring Boot,高德地图,DeepSeek"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    make_document()
